"""
Tests for the dynamic-docx-fill layer: content-control structure discovery,
fill, and read-back verification.

There is no real Word-authored content-control fixture in this repo, so the
fixture here is built with genuine OOXML w:sdt elements (the same XML Word
itself writes for a text/checkbox/dropdown content control), not a mock -
the code under test never knows the difference. This exercises the real
`renderers.docx_structure`/`renderers.docx_filler` modules end to end.
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_ROOT))

from renderers import docx_structure  # noqa: E402
from renderers.docx_filler import fill_docx_controls, read_back_docx_controls  # noqa: E402


def _sdt_text(tag: str, alias: str | None = None, placeholder: str = "Click here to enter text.") -> object:
    alias_xml = f'<w:alias w:val="{alias}"/>' if alias else ""
    xml = f"""
    <w:sdt {nsdecls("w")}>
      <w:sdtPr>
        {alias_xml}
        <w:tag w:val="{tag}"/>
        <w:id w:val="1"/>
        <w:showingPlcHdr/>
        <w:text/>
      </w:sdtPr>
      <w:sdtContent>
        <w:r><w:t>{placeholder}</w:t></w:r>
      </w:sdtContent>
    </w:sdt>
    """
    return parse_xml(xml)


def _sdt_checkbox(tag: str, alias: str, checked: bool = False) -> object:
    val = "1" if checked else "0"
    glyph = chr(0x2612) if checked else chr(0x2610)
    xml = f"""
    <w:sdt {nsdecls("w", "w14")}>
      <w:sdtPr>
        <w:alias w:val="{alias}"/>
        <w:tag w:val="{tag}"/>
        <w:id w:val="2"/>
        <w14:checkbox>
          <w14:checked w14:val="{val}"/>
          <w14:checkedState w14:val="2612" w14:font="MS Gothic"/>
          <w14:uncheckedState w14:val="2610" w14:font="MS Gothic"/>
        </w14:checkbox>
      </w:sdtPr>
      <w:sdtContent>
        <w:r><w:rPr><w:rFonts w:ascii="MS Gothic" w:hAnsi="MS Gothic"/></w:rPr><w:t>{glyph}</w:t></w:r>
      </w:sdtContent>
    </w:sdt>
    """
    return parse_xml(xml)


def _sdt_dropdown(tag: str, alias: str, items: list[tuple[str, str]], placeholder: str = "Choose an item.") -> object:
    list_items = "".join(f'<w:listItem w:displayText="{d}" w:value="{v}"/>' for d, v in items)
    xml = f"""
    <w:sdt {nsdecls("w")}>
      <w:sdtPr>
        <w:alias w:val="{alias}"/>
        <w:tag w:val="{tag}"/>
        <w:id w:val="3"/>
        <w:showingPlcHdr/>
        <w:dropDownList>{list_items}</w:dropDownList>
      </w:sdtPr>
      <w:sdtContent>
        <w:r><w:t>{placeholder}</w:t></w:r>
      </w:sdtContent>
    </w:sdt>
    """
    return parse_xml(xml)


def build_fixture_docx() -> bytes:
    doc = Document()

    p1 = doc.add_paragraph("Full Name:")
    p1._p.append(_sdt_text("full_name", alias="Full Name"))

    # No alias set here on purpose - exercises the _context_label fallback
    # (nearest inline text before the control), not the alias shortcut.
    p2 = doc.add_paragraph("Date of Loss:")
    p2._p.append(_sdt_text("loss_date"))

    p3 = doc.add_paragraph("Are you the policyholder?")
    p3._p.append(_sdt_checkbox("is_policyholder", alias="Is Policyholder", checked=False))

    p4 = doc.add_paragraph("State:")
    p4._p.append(_sdt_dropdown("state", alias="State", items=[
        ("Maharashtra", "MH"), ("Karnataka", "KA"), ("Tamil Nadu", "TN"),
    ]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture(scope="module")
def fixture_bytes() -> bytes:
    return build_fixture_docx()


@pytest.fixture(scope="module")
def draft(fixture_bytes) -> dict:
    return docx_structure.build_draft(fixture_bytes)


# =============================================================================
# Structure discovery
# =============================================================================

def test_discovers_all_control_types(draft):
    assert draft["stats"]["controls"] == 4
    types = {c["name"]: c["type"] for c in draft["controls"]}
    assert types == {
        "full_name": "text",
        "loss_date": "text",
        "is_policyholder": "checkbox",
        "state": "dropdown",
    }


def test_alias_is_used_as_label_when_present(draft):
    by_name = {c["name"]: c for c in draft["controls"]}
    assert by_name["full_name"]["label"] == "Full Name"
    assert by_name["is_policyholder"]["label"] == "Is Policyholder"


def test_context_label_fallback_when_no_alias(draft):
    # loss_date has no <w:alias> - label must come from the inline text
    # ("Date of Loss:") preceding the control in the same paragraph.
    by_name = {c["name"]: c for c in draft["controls"]}
    assert by_name["loss_date"]["alias"] is None
    assert by_name["loss_date"]["label"] == "Date of Loss"


def test_dropdown_choices_are_harvested(draft):
    by_name = {c["name"]: c for c in draft["controls"]}
    displays = {c["display"] for c in by_name["state"]["choices"]}
    assert displays == {"Maharashtra", "Karnataka", "Tamil Nadu"}


def test_placeholder_controls_report_empty_current_text(draft):
    by_name = {c["name"]: c for c in draft["controls"]}
    assert by_name["full_name"]["is_placeholder"] is True
    assert by_name["full_name"]["current_text"] == ""


def test_checkbox_reports_its_current_checked_state(draft):
    by_name = {c["name"]: c for c in draft["controls"]}
    assert by_name["is_policyholder"]["checked"] is False


def test_meaning_is_left_for_the_agent_not_decided_here(draft):
    assert draft["status"] == "DRAFT_STRUCTURAL"
    assert all(c["concept"] is None and c["dtype"] is None for c in draft["controls"])


def test_fingerprint_is_stable_for_the_same_docx(fixture_bytes):
    a = docx_structure.build_draft(fixture_bytes)["fingerprint"]
    b = docx_structure.build_draft(fixture_bytes)["fingerprint"]
    assert a == b and len(a) == 16


def test_docx_with_no_content_controls_reports_zero():
    plain = Document()
    plain.add_paragraph("Just prose, no form fields here.")
    buf = io.BytesIO()
    plain.save(buf)
    draft = docx_structure.build_draft(buf.getvalue())
    assert draft["stats"]["controls"] == 0


# =============================================================================
# Fill + verify round trip
# =============================================================================

def test_fill_then_read_back_matches_for_every_control_type(fixture_bytes):
    values = {"full_name": "Aarav Sharma", "loss_date": "12/03/2025"}
    checks = {"is_policyholder": True}
    choices = {"state": "Karnataka"}

    filled = fill_docx_controls(fixture_bytes, values, checks, choices)
    back = read_back_docx_controls(filled)

    assert back["full_name"] == "Aarav Sharma"
    assert back["loss_date"] == "12/03/2025"
    assert back["is_policyholder"] is True
    assert back["state"] == "Karnataka"


def test_fill_clears_the_placeholder_flag(fixture_bytes):
    filled = fill_docx_controls(fixture_bytes, {"full_name": "Priya Iyer"})
    draft = docx_structure.build_draft(filled)
    by_name = {c["name"]: c for c in draft["controls"]}
    assert by_name["full_name"]["is_placeholder"] is False


def test_checkbox_uncheck_round_trips_too(fixture_bytes):
    filled = fill_docx_controls(fixture_bytes, checks={"is_policyholder": False})
    back = read_back_docx_controls(filled)
    assert back["is_policyholder"] is False


def test_dropdown_rejects_an_option_the_control_does_not_offer(fixture_bytes):
    with pytest.raises(ValueError, match="not one of the control's listItems"):
        fill_docx_controls(fixture_bytes, choices={"state": "Delhi"})


def test_fill_ignores_unknown_control_names_rather_than_raising(fixture_bytes):
    # Mirrors the PDF tool layer's "no matching widgets" -> empty result
    # philosophy: an agent passing a stale/misremembered name should not
    # blow up the whole fill, just not touch anything.
    filled = fill_docx_controls(fixture_bytes, values={"does_not_exist": "x"})
    back = read_back_docx_controls(filled)
    assert back["full_name"] == ""
