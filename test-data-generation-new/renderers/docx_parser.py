import io
from docx import Document


def extract_docx_layout(docx_bytes: bytes) -> dict:
    doc = Document(io.BytesIO(docx_bytes))
    layout = {
        "file_type": "docx",
        "headings": [],
        "field_labels": [],
        "paragraphs_sample": [],
        "tables": [],
    }

    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()
        if not text:
            continue
        if style.startswith("Heading"):
            layout["headings"].append({"level": style, "text": text})
        elif text.endswith(":") and len(text) < 60:
            layout["field_labels"].append(text)
        elif len(layout["paragraphs_sample"]) < 10:
            layout["paragraphs_sample"].append(text[:300])

    for table in doc.tables:
        headers = []
        if table.rows:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
        layout["tables"].append({
            "rows": len(table.rows),
            "cols": len(table.columns),
            "headers": headers,
        })

    return layout
